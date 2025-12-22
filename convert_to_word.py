"""
Script para converter technical_manual.md para Word (.docx)
Requer: pip install python-docx markdown
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_word(md_file, docx_file):
    """Converte Markdown para Word com formatação básica."""
    
    # Lê o arquivo Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Cria documento Word
    doc = Document()
    
    # Configurações de estilo
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Processa linha por linha
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Ignora linhas vazias múltiplas
        if not line.strip():
            i += 1
            continue
        
        # Título H1 (# )
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Título H2 (## )
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
        
        # Título H3 (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
        
        # Título H4 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
        
        # Separador (---)
        elif line.strip() == '---':
            doc.add_paragraph()  # Apenas adiciona espaço
        
        # Bloco de código (```)
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Adiciona código com formatação
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Adiciona fundo cinza claro (simulado com borda)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.25)
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)
        
        # Lista não ordenada (- ou *)
        elif re.match(r'^[\-\*]\s+', line):
            text = re.sub(r'^[\-\*]\s+', '', line)
            text = clean_markdown_formatting(text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Lista ordenada (1. )
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            text = clean_markdown_formatting(text)
            doc.add_paragraph(text, style='List Number')
        
        # Tabela (| ... |)
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Volta uma linha
            
            # Processa tabela
            rows = []
            for tline in table_lines:
                if '---' in tline:  # Linha separadora
                    continue
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                rows.append(cells)
            
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = clean_markdown_formatting(cell_data)
        
        # Parágrafo normal
        else:
            text = clean_markdown_formatting(line)
            if text.strip():
                doc.add_paragraph(text)
        
        i += 1
    
    # Salva documento
    doc.save(docx_file)
    print(f"✅ Documento Word gerado: {docx_file}")

def clean_markdown_formatting(text):
    """Remove formatação Markdown básica."""
    # Remove ** (negrito)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove ` (código inline)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove links [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text

if __name__ == '__main__':
    script_dir = Path(__file__).parent
    md_file = script_dir / 'technical_manual.md'
    docx_file = script_dir / 'technical_manual.docx'
    
    print(f"📄 Convertendo {md_file.name} para Word...")
    markdown_to_word(md_file, docx_file)
    print(f"✅ Conversão concluída!")
    print(f"📁 Arquivo salvo em: {docx_file}")
